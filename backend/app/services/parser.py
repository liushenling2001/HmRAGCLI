from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
import re
import subprocess
import sys
from zipfile import ZipFile

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader

from app.core.config import settings
from app.services import converter as converter_service


@dataclass
class ParsedChunk:
    content: str
    chunk_type: str = "text"
    title: str | None = None
    page_no: int | None = None
    metadata: dict | None = None


@dataclass
class ParsePlan:
    original_path: str
    resolved_path: Path
    requested_mode: str
    selected_mode: str
    backend: str
    conversion: dict


@dataclass
class PreparedParseResult:
    preview_text: str
    parse_metadata: dict
    parsed_chunks: list[ParsedChunk]


class EnhancedParserNotReadyError(RuntimeError):
    pass


def _safe_path(path_str: str) -> Path:
    return Path(path_str)


def _split_text(text: str, max_chars: int = 1800) -> list[str]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [block.strip() for block in normalized.split("\n\n") if block.strip()]
    chunks: list[str] = []
    current = ""
    for block in blocks:
        if len(current) + len(block) + 2 <= max_chars:
            current = f"{current}\n\n{block}".strip()
        else:
            if current:
                chunks.append(current)
            if len(block) <= max_chars:
                current = block
            else:
                start = 0
                while start < len(block):
                    chunks.append(block[start : start + max_chars])
                    start += max_chars
                current = ""
    if current:
        chunks.append(current)
    return chunks


def _paragraph_style_name(paragraph: Paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip()
    except Exception:
        return ""


def _heading_level(paragraph: Paragraph, text: str) -> int | None:
    style_name = _paragraph_style_name(paragraph)
    lowered = style_name.lower()
    match = re.search(r"heading\s*(\d+)", lowered)
    if match:
        return int(match.group(1))
    match = re.search(r"标题\s*(\d+)", style_name)
    if match:
        return int(match.group(1))
    if re.match(r"^第[一二三四五六七八九十百千万0-9]+[章节部分篇]\s*", text):
        return 1
    if re.match(r"^\d+(\.\d+){0,3}\s+", text):
        depth = text.split(" ", 1)[0].count(".") + 1
        return min(depth, 6)
    return None


def _is_list_paragraph(paragraph: Paragraph, text: str) -> bool:
    style_name = _paragraph_style_name(paragraph).lower()
    if any(token in style_name for token in ("list", "bullet", "number", "列表")):
        return True
    return bool(re.match(r"^([\-*•]|[0-9]+[\.、)]|[一二三四五六七八九十]+、)\s+", text))


def _make_heading_path(path_items: list[str]) -> str | None:
    if not path_items:
        return None
    return " / ".join(path_items)


def _parse_docx_paragraphs(doc: DocxDocument) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    heading_stack: list[str] = []
    current_lines: list[str] = []
    current_title: str | None = None
    current_meta: dict = {
        "block_role": "paragraph",
        "paragraph_styles": [],
        "heading_path": None,
        "heading_level": None,
    }

    def flush_current() -> None:
        nonlocal current_lines, current_title, current_meta
        text = "\n".join(line for line in current_lines if line.strip()).strip()
        if not text:
            current_lines = []
            current_title = None
            current_meta = {
                "block_role": "paragraph",
                "paragraph_styles": [],
                "heading_path": _make_heading_path(heading_stack),
                "heading_level": None,
            }
            return
        split_chunks = _split_text(text)
        for idx, chunk_text in enumerate(split_chunks):
            chunk_meta = {
                **current_meta,
                "heading_path": current_meta.get("heading_path") or _make_heading_path(heading_stack),
                "is_continuation": idx > 0,
            }
            chunks.append(
                ParsedChunk(
                    content=chunk_text,
                    chunk_type="text",
                    title=current_title,
                    metadata=chunk_meta,
                )
            )
        current_lines = []
        current_title = None
        current_meta = {
            "block_role": "paragraph",
            "paragraph_styles": [],
            "heading_path": _make_heading_path(heading_stack),
            "heading_level": None,
        }

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            flush_current()
            continue

        style_name = _paragraph_style_name(paragraph)
        heading_level = _heading_level(paragraph, text)
        is_list = _is_list_paragraph(paragraph, text)

        if heading_level is not None:
            flush_current()
            while len(heading_stack) >= heading_level:
                heading_stack.pop()
            heading_stack.append(text)
            chunks.append(
                ParsedChunk(
                    content=text,
                    chunk_type="heading",
                    title=text,
                    metadata={
                        "block_role": "heading",
                        "search_role": "heading",
                        "search_title": text,
                        "heading_level": heading_level,
                        "heading_path": _make_heading_path(heading_stack),
                        "paragraph_style": style_name,
                    },
                )
            )
            continue

        target_role = "list" if is_list else "paragraph"
        if current_lines and current_meta.get("block_role") != target_role:
            flush_current()

        if not current_lines:
            current_title = heading_stack[-1] if heading_stack else None
            current_meta = {
                "block_role": target_role,
                "search_role": "body" if target_role == "paragraph" else "list",
                "paragraph_styles": [style_name] if style_name else [],
                "heading_path": _make_heading_path(heading_stack),
                "heading_level": len(heading_stack) or None,
            }
        else:
            if style_name and style_name not in current_meta["paragraph_styles"]:
                current_meta["paragraph_styles"].append(style_name)

        current_lines.append(text)
        if sum(len(line) for line in current_lines) >= 1800:
            flush_current()

    flush_current()
    return chunks


def _parse_docx_tables(doc: DocxDocument) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    for table_no, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        headers: list[str] = []
        row_count = 0
        col_count = 0
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if values:
                row_count += 1
                col_count = max(col_count, len(values))
                if not headers:
                    headers = values
                rows.append(" | ".join(values))
        if not rows:
            continue
        table_text = f"Table {table_no}\n" + "\n".join(rows)
        for chunk in _split_text(table_text):
            chunks.append(
                ParsedChunk(
                    content=chunk,
                    chunk_type="table",
                    title=f"Table {table_no}",
                    metadata={
                        "block_role": "table",
                        "search_role": "table",
                        "search_title": f"Table {table_no}",
                        "table_no": table_no,
                        "row_count": row_count,
                        "col_count": col_count,
                        "headers": headers[:20],
                    },
                )
            )
    return chunks


def _parse_docx_images(path: Path) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    try:
        with ZipFile(path) as archive:
            media_names = sorted(
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            )
    except Exception:
        return chunks

    for image_no, media_name in enumerate(media_names, start=1):
        image_name = Path(media_name).name
        chunks.append(
            ParsedChunk(
                content=f"Word image placeholder: {image_name}",
                chunk_type="figure",
                title=f"Image {image_no}",
                metadata={
                    "block_role": "figure",
                    "search_role": "figure",
                    "search_title": f"Image {image_no}",
                    "image_no": image_no,
                    "image_name": image_name,
                },
            )
        )
    return chunks


def _parse_path_light(path: Path) -> list[ParsedChunk]:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [ParsedChunk(content=chunk) for chunk in _split_text(text)]
    if suffix == ".docx":
        doc = DocxDocument(str(path))
        chunks: list[ParsedChunk] = []
        chunks.extend(_parse_docx_paragraphs(doc))
        chunks.extend(_parse_docx_tables(doc))
        chunks.extend(_parse_docx_images(path))
        return chunks
    if suffix == ".pdf":
        return _parse_pdf_with_timeout(path)
    if suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
        wb = load_workbook(path, read_only=True, data_only=True)
        chunks: list[ParsedChunk] = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if values:
                    rows.append(" | ".join(values))
            if rows:
                text = f"Sheet: {sheet.title}\n" + "\n".join(rows[:200])
                for chunk in _split_text(text):
                    chunks.append(
                        ParsedChunk(
                            content=chunk,
                            chunk_type="sheet",
                            title=sheet.title,
                            metadata={
                                "block_role": "sheet",
                                "search_role": "sheet",
                                "search_title": sheet.title,
                                "sheet_name": sheet.title,
                            },
                        )
                    )
        return chunks
    return []


def _parse_pdf_light(path_str: str) -> list[ParsedChunk]:
    reader = PdfReader(path_str)
    chunks: list[ParsedChunk] = []
    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        for chunk in _split_text(text):
            chunks.append(ParsedChunk(content=chunk, page_no=page_no))
    return chunks


def _parse_pdf_with_timeout(path: Path) -> list[ParsedChunk]:
    timeout_seconds = max(1, int(settings.pdf_parse_timeout_seconds))
    try:
        script = (
            "import json, sys\n"
            "from app.services.parser import _parse_pdf_light\n"
            "chunks = _parse_pdf_light(sys.argv[1])\n"
            "print(json.dumps([chunk.__dict__ for chunk in chunks], ensure_ascii=True))\n"
        )
        completed = subprocess.run(
            [sys.executable, "-c", script, str(path)],
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
    except subprocess.TimeoutExpired as exc:
        raise TimeoutError(
            f"PDF parsing timed out after {timeout_seconds}s: {path.name}"
        ) from exc
    if completed.returncode != 0:
        message = completed.stderr.strip() or completed.stdout.strip() or "PDF parsing failed"
        raise RuntimeError(message)
    payload = json.loads(completed.stdout or "[]")
    return [ParsedChunk(**item) for item in payload]


def _build_pdf_preview(path: Path, max_chars: int) -> str:
    reader = PdfReader(str(path))
    page_limit = max(1, settings.pdf_preview_page_limit)
    parts: list[str] = []
    total = 0
    for page_no, page in enumerate(reader.pages, start=1):
        if page_no > page_limit or total >= max_chars:
            break
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        clipped = text[: max_chars - total]
        parts.append(clipped)
        total += len(clipped)
    return "\n\n".join(parts)[:max_chars]


def _build_docx_preview(path: Path, max_chars: int) -> str:
    chunks = _parse_path_light(path)
    if not chunks:
        return ""
    preview = "\n\n".join(chunk.content for chunk in chunks[:6])
    return preview[:max_chars]


def _choose_parse_mode(path: Path, requested_mode: str) -> tuple[str, str]:
    requested = (requested_mode or settings.parse_mode or "auto").strip().lower()
    backend = (settings.enhanced_parser_backend or "none").strip().lower()
    suffix = path.suffix.lower()
    size_kb = (path.stat().st_size // 1024) if path.exists() else 0

    if requested == "light":
        return "light", "builtin-light"
    if requested == "enhanced":
        return "enhanced", backend

    if backend not in {"docling", "mineru"}:
        return "light", "builtin-light"

    if suffix == ".pdf" and size_kb >= settings.enhanced_pdf_size_kb:
        return "enhanced", backend
    if suffix == ".docx" and size_kb >= settings.enhanced_docx_size_kb:
        return "enhanced", backend
    return "light", "builtin-light"


def resolve_parse_plan(path_str: str, mode: str | None = None) -> ParsePlan:
    resolved_path, conversion = converter_service.resolve_input_path(path_str)
    selected_mode, backend = _choose_parse_mode(resolved_path, mode or settings.parse_mode)
    return ParsePlan(
        original_path=path_str,
        resolved_path=resolved_path,
        requested_mode=(mode or settings.parse_mode or "auto").strip().lower(),
        selected_mode=selected_mode,
        backend=backend,
        conversion=conversion,
    )


def _parse_path_enhanced(plan: ParsePlan) -> list[ParsedChunk]:
    backend = plan.backend
    if backend == "none":
        raise EnhancedParserNotReadyError("Enhanced parser requested but no backend is configured")
    if backend == "docling":
        from app.services import docling_parser

        try:
            return docling_parser.parse_with_docling(plan.resolved_path)
        except docling_parser.DoclingUnavailableError as exc:
            raise EnhancedParserNotReadyError(str(exc)) from exc
    if backend == "mineru":
        raise EnhancedParserNotReadyError(
            "MinerU enhanced parser is not wired into this offline build yet. Prepare dependencies offline before enabling it."
        )
    raise EnhancedParserNotReadyError(f"Unsupported enhanced parser backend: {backend}")


def parse_path(path_str: str, mode: str | None = None) -> list[ParsedChunk]:
    plan = resolve_parse_plan(path_str, mode=mode)
    if plan.selected_mode == "enhanced":
        return _parse_path_enhanced(plan)
    return _parse_path_light(plan.resolved_path)


def build_preview_text(path_str: str, max_chars: int = 2000, mode: str | None = None) -> str:
    plan = resolve_parse_plan(path_str, mode=mode)
    suffix = plan.resolved_path.suffix.lower()
    if suffix == ".pdf":
        return _build_pdf_preview(plan.resolved_path, max_chars)
    if suffix in {".docx", ".doc"}:
        return _build_docx_preview(plan.resolved_path, max_chars)
    chunks = _parse_path_light(plan.resolved_path)
    if not chunks:
        return ""
    preview = "\n\n".join(chunk.content for chunk in chunks[:3])
    return preview[:max_chars]


def build_parse_metadata(path_str: str, mode: str | None = None) -> dict:
    plan = resolve_parse_plan(path_str, mode=mode)
    return {
        **plan.conversion,
        "requested_mode": plan.requested_mode,
        "selected_mode": plan.selected_mode,
        "parser_backend": plan.backend,
    }


def _preview_from_parsed_chunks(chunks: list[ParsedChunk], max_chars: int = 2000) -> str:
    if not chunks:
        return ""
    parts: list[str] = []
    total = 0
    for chunk in chunks:
        text = (chunk.content or "").strip()
        if not text:
            continue
        remaining = max_chars - total
        if remaining <= 0:
            break
        clipped = text[:remaining]
        parts.append(clipped)
        total += len(clipped)
    return "\n\n".join(parts)[:max_chars]


def prepare_parse_result(path_str: str, mode: str | None = None, *, max_preview_chars: int = 2000) -> PreparedParseResult:
    plan = resolve_parse_plan(path_str, mode=mode)
    parse_metadata = {
        **plan.conversion,
        "requested_mode": plan.requested_mode,
        "selected_mode": plan.selected_mode,
        "parser_backend": plan.backend,
    }
    if plan.selected_mode == "enhanced":
        parsed_chunks = _parse_path_enhanced(plan)
    else:
        parsed_chunks = _parse_path_light(plan.resolved_path)
    return PreparedParseResult(
        preview_text=_preview_from_parsed_chunks(parsed_chunks, max_chars=max_preview_chars),
        parse_metadata=parse_metadata,
        parsed_chunks=parsed_chunks,
    )
